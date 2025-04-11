from docx import Document
import re
from docx.shared import Pt

def load_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def identify_elements(text):
    headings = re.findall(r'^[A-Z].*$', text, re.MULTILINE)  # Пример заголовков
    lists = re.findall(r'^\s*[-•]\s+.*$', text, re.MULTILINE)  # Пример списков
    return headings, lists

def format_document(file_path):
    doc = Document(file_path)
    for para in doc.paragraphs:
        if para.text.startswith('Глава'):
            para.style.font.size = Pt(14)  # Применяем шрифт для заголовков
    
    # Сохраняем документ после форматирования
    doc.save('formatted_document.docx')
